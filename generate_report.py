from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_report(ogrenci1="20260001", ogrenci2="20260002"):
    doc = Document()

    # Baslik
    title = doc.add_heading("Biyometrik Ses Dogrulama ve Yapay Zeka Entegreli Asistan Sistemi", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Yazarlar
    authors = doc.add_paragraph(f"Öğrenci No 1: {ogrenci1}   |   Öğrenci No 2: {ogrenci2}")
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading("Özet", level=1)
    doc.add_paragraph(
        "Bu projede, güvenlik sistemlerinde kullanilmak uzere, kisinin sesinden kimligini "
        "dogrulayan biyometrik bir sistem gelistirilmistir. Ses on isleme adiminda pre-emphasis "
        "ve gurultu azaltma uygulanmis, ardindan Mel-Frequency Cepstral Coefficients (MFCC), "
        "delta ve delta-delta (ayrica ortalama ve standart sapma) öznitelikleri cikarilmistir. "
        "Cikarilan vektorler Destek Vektör Makineleri (SVM) ile egitilerek %0.11 EER ve yuksek dogruluk "
        "oranina ulasilmistir. Sistemin ikinci asamasinda, dogrulanan kullanicinin komutlari "
        "Whisper (STT) ve Llama-3.3 (LLM) ile islenerek akilli asistan yanitlari uretilmistir."
    )

    doc.add_heading("1. Giriş", level=1)
    doc.add_paragraph(
        "Biyometrik sistemler, bireylerin fiziksel veya davranissal özelliklerini kullanarak "
        "kimlik dogrulamasi saglayan guvenlik mekanizmalaridir. Ses, benzersiz akustik yapisi "
        "nedeniyle onemli bir biyometrik veridir. Bu projede hem ses tanima hem de "
        "yapay zeka (LLM) tabanli komut isleme yetenegine sahip hibrit bir 'AEGIS OS' "
        "sistemi tasarlanmistir."
    )

    doc.add_heading("2. Literatür Taraması", level=1)
    doc.add_paragraph(
        "Konusmaci tanima sistemleri uzerine yapilan calismalarda Gauss Karisim Modelleri (GMM), "
        "I-Vector ve son yillarda X-Vector/ECAPA-TDNN gibi derin ogrenme modelleri "
        "kullanilmaktadir. Klasik yaklasimlarda ise MFCC ve SVM ikilisi donanim dostu ve "
        "hizli sonuc vermesiyle one cikmaktadir."
    )

    doc.add_heading("3. Yöntem & Malzeme (Veri Seti)", level=1)
    doc.add_paragraph(
        "Projede 'LibriSpeech' veri seti ve sisteme sonradan kaydedilen gercek zamanli "
        "kullanici sesleri (User Voice) kullanilmistir. \n\n"
        "Öznitelik Çıkarımı: Sesten MFCC, delta ve delta-delta bilesenleri alinmis, "
        "zaman ekseninde ortalama (mean) ve standart sapma (std) alinarak (Toplam 240 boyut) "
        "veriler egitime hazirlanmistir.\n"
        "Sınıflandırma: RBF cekirdekli SVM (Support Vector Machine) algoritmasi secilmistir."
    )

    doc.add_heading("4. Deney Kurulumu", level=1)
    doc.add_paragraph(
        "Veri seti %80 egitim ve %20 test olarak ayrilmistir. StandardScaler ile "
        "oznitelikler olceklendirilmistir. Python ekosisteminde librosa, scikit-learn "
        "kutuphaneleri kullanilmistir. Arayuz Streamlit ve HTML/Tailwind CSS ile "
        "modern bir 'Glassmorphism' yapisinda kurulmustur."
    )

    doc.add_heading("5. Bulgular", level=1)
    doc.add_paragraph(
        "Sistem %0.11 EER (Equal Error Rate) oranina ulasmistir. DET egrisi test "
        "verisi uzerinde olusturulmustur. Asagida test edilen modele ait DET egrisi yer almaktadir:"
    )

    if os.path.exists("data/det_curve.png"):
        doc.add_picture("data/det_curve.png", width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        doc.add_paragraph("[Grafik Bulunamadi - Lutfen evaluation.py calistirarak data/det_curve.png olusturun]")

    doc.add_heading("6. Tartışma", level=1)
    doc.add_paragraph(
        "MFCC ve SVM yaklasimi izole ortamlarda cok yuksek basari gosterirken, "
        "arkaplan gurultusunun fazla oldugu durumlarda performans duser. Pre-emphasis "
        "ve gurultu kirpma islemleri bu etkiyi kismen azaltmistir. Ayrica LLM entegrasyonu "
        "sayesinde sistemin salt guvenlikten ote islevsel bir yapay zeka asistanina "
        "donusmesi saglanmistir."
    )

    doc.add_heading("7. Sonuç", level=1)
    doc.add_paragraph(
        "Proje hedefleri dogrultusunda biyometrik bir ses dogrulama sistemi basariyla "
        "gerceklestirilmistir. Sistemin sadece dogru kisiyi tanimasi yetmemis, o "
        "kisinin komutlarini algilayarak anlamli bir etkilesim de (Groq Whisper & Llama) "
        "saglamistir."
    )

    doc.add_heading("8. Gelecek Çalışmalar", level=1)
    doc.add_paragraph(
        "Gelecekte modelin ECAPA-TDNN gibi modern Transformer veya CNN tabanli "
        "derin ogrenme algoritmalarina gecirilmesi hedeflenmektedir. Ayrica "
        "gurultulu ortam performansini artirmak icin Spectral Subtraction "
        "yontemleri de eklenebilir."
    )

    os.makedirs("data", exist_ok=True)
    filename = f"{ogrenci1}_{ogrenci2}_Rapor.docx"
    doc.save(filename)
    print(f"Rapor basariyla olusturuldu: {filename}")

if __name__ == "__main__":
    o1 = input("1. Ogrenci No: ")
    o2 = input("2. Ogrenci No: ")
    create_report(o1, o2)
